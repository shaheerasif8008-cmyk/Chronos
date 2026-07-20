"""Durable DOCX and PDF exports for completed research reports."""
from __future__ import annotations

import asyncio
from contextlib import asynccontextmanager
from datetime import datetime, timezone
import hashlib
from html import escape as html_escape
from io import BytesIO
import re
import textwrap
from typing import Any, Literal

from sqlalchemy import select, text, update

from core.artifacts import (
    get_artifact,
    read_artifact_content,
    save_artifact,
    set_artifact_project,
)
from core.db import engine, reflect_table


ResearchExportFormat = Literal["docx", "pdf"]
_KINDS = {"docx": "research_report_docx", "pdf": "research_report_pdf"}
_MIMES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}
_LOCK_WAIT_SECONDS = 30.0
_MAX_REPORT_BYTES = 5 * 1024 * 1024
_MAX_CITATION_BYTES = 5 * 1024 * 1024
_MAX_CITATIONS = 2_000


class ResearchExportError(ValueError):
    """A research run cannot be exported in its current state."""


def _export_lock_key(org_id: str, report_id: str, export_format: str) -> int:
    """Stable signed bigint key for a Postgres session advisory lock."""
    digest = hashlib.blake2b(
        f"chronos:research-export:{org_id}:{report_id}:{export_format}".encode("utf-8"),
        digest_size=8,
    ).digest()
    return int.from_bytes(digest, "big", signed=True)


@asynccontextmanager
async def _export_lock(org_id: str, report_id: str, export_format: str):
    """Serialize one report/format export across every API replica.

    A session lock is intentional: the artifact insert is committed while the
    same connection still holds this lock, so waiters cannot observe a partial
    export. Closing the connection also releases the lock if cancellation lands
    while the explicit unlock is running.
    """
    key = _export_lock_key(org_id, report_id, export_format)
    async with engine.connect() as conn:
        deadline = asyncio.get_running_loop().time() + _LOCK_WAIT_SECONDS
        acquired = False
        while not acquired:
            acquired = bool(
                await conn.scalar(
                    text("SELECT pg_try_advisory_lock(:key)"), {"key": key}
                )
            )
            if acquired:
                break
            if asyncio.get_running_loop().time() >= deadline:
                raise ResearchExportError(
                    "This report export is already being created. Try again shortly."
                )
            await asyncio.sleep(0.05)
        try:
            yield conn
        finally:
            if acquired:
                try:
                    await conn.execute(
                        text("SELECT pg_advisory_unlock(:key)"), {"key": key}
                    )
                except Exception:
                    # Connection close is the fail-safe lock release path.
                    pass


async def _find_existing_export(
    artifacts: Any,
    *,
    org_id: str,
    report_id: str,
    export_format: str,
    conn: Any | None = None,
) -> dict[str, Any] | None:
    query = (
        select(artifacts)
        .where(
            artifacts.c.organization_id == org_id,
            artifacts.c.parent_artifact_id == report_id,
            artifacts.c.kind == _KINDS[export_format],
            artifacts.c.is_deleted == False,  # noqa: E712
        )
        .order_by(artifacts.c.created_at.desc())
        .limit(1)
    )
    if conn is not None:
        existing = (await conn.execute(query)).mappings().first()
    else:
        async with engine.begin() as managed_conn:
            existing = (await managed_conn.execute(query)).mappings().first()
    return dict(existing) if existing else None


async def _set_project_and_fetch(
    artifacts: Any,
    artifact_id: str,
    *,
    project_id: str | None,
    org_id: str,
    conn: Any,
) -> dict[str, Any] | None:
    if project_id:
        await conn.execute(
            update(artifacts)
            .where(
                artifacts.c.id == artifact_id,
                artifacts.c.organization_id == org_id,
                artifacts.c.is_deleted == False,  # noqa: E712
            )
            .values(project_id=project_id, updated_at=datetime.now(timezone.utc))
        )
    row = (
        await conn.execute(
            select(artifacts).where(
                artifacts.c.id == artifact_id,
                artifacts.c.organization_id == org_id,
            )
        )
    ).mappings().first()
    return dict(row) if row else None


def _clean_title(question: str, extension: str) -> str:
    title = re.sub(r"[\r\n\t]+", " ", question).strip()
    title = re.sub(r"\s+", " ", title)[:140] or "Research report"
    return f"{title}.{extension}"


def _report_body(markdown: str) -> str:
    """Remove generated Sources/Limitations sections before canonical appendices."""
    lines = markdown.replace("\r\n", "\n").split("\n")
    output: list[str] = []
    skipping = False
    skip_level = 0
    for line in lines:
        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if heading:
            level = len(heading.group(1))
            label = re.sub(r"[*_`]", "", heading.group(2)).strip().lower().rstrip(":")
            if label in {"sources", "source", "citations", "references", "limitations"}:
                skipping = True
                skip_level = level
                continue
            if skipping and level <= skip_level:
                skipping = False
        if not skipping:
            output.append(line)
    return "\n".join(output).strip()


def _markdown_blocks(markdown: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    paragraph: list[str] = []
    code: list[str] = []
    in_code = False

    def flush_paragraph() -> None:
        if paragraph:
            blocks.append(("paragraph", " ".join(part.strip() for part in paragraph).strip()))
            paragraph.clear()

    for line in markdown.splitlines():
        if line.strip().startswith("```"):
            if in_code:
                blocks.append(("code", "\n".join(code)))
                code.clear()
                in_code = False
            else:
                flush_paragraph()
                in_code = True
            continue
        if in_code:
            code.append(line)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            blocks.append((f"heading{min(len(heading.group(1)), 3)}", heading.group(2).strip()))
            continue
        bullet = re.match(r"^\s*[-*+]\s+(.+)$", line)
        numbered = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if bullet or numbered:
            flush_paragraph()
            blocks.append(("bullet" if bullet else "number", (bullet or numbered).group(1).strip()))
            continue
        if not line.strip():
            flush_paragraph()
        else:
            paragraph.append(line)
    flush_paragraph()
    if code:
        blocks.append(("code", "\n".join(code)))
    return [(kind, text) for kind, text in blocks if text]


def _plain_markdown(text: str) -> str:
    text = re.sub(r"!\[([^]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    # Preserve underscores because they are common in identifiers and URLs;
    # deleting them corrupts evidence (for example, release_sha).
    text = re.sub(r"(?<!\\)[*~`]", "", text)
    return text.strip()


def _add_page_field(paragraph: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def build_research_docx(run: dict[str, Any], report: str, citations: list[dict[str, Any]]) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, RGBColor(46, 116, 181)),
        ("Heading 2", 13, 12, 6, RGBColor(46, 116, 181)),
        ("Heading 3", 12, 8, 4, RGBColor(31, 77, 120)),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(92)
    title.paragraph_format.space_after = Pt(18)
    title_run = title.add_run("Research report")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    question = document.add_paragraph()
    question.paragraph_format.space_after = Pt(24)
    question_run = question.add_run(str(run.get("question") or "Untitled research"))
    question_run.font.size = Pt(16)
    question_run.font.color.rgb = RGBColor(45, 64, 89)
    metadata = document.add_paragraph()
    metadata.add_run(f"Depth: {run.get('depth') or 'standard'}\n").bold = True
    metadata.add_run(f"Completed: {run.get('completed_at') or datetime.now(timezone.utc).isoformat()}\n")
    metadata.add_run(f"Citation policy: {run.get('citation_policy') or 'required'}\n")
    metadata.add_run(f"Sources cited: {len(citations)}")
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    body = _report_body(report)
    for kind, raw_text in _markdown_blocks(body):
        text = raw_text if kind == "code" else _plain_markdown(raw_text)
        if not text:
            continue
        if kind.startswith("heading"):
            document.add_heading(text, level=int(kind[-1]))
        elif kind == "bullet":
            document.add_paragraph(text, style="List Bullet")
        elif kind == "number":
            document.add_paragraph(text, style="List Number")
        elif kind == "code":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(8)
            run_code = paragraph.add_run(text[:50_000])
            run_code.font.name = "Courier New"
            run_code.font.size = Pt(9)
        else:
            document.add_paragraph(text)

    document.add_heading("Limitations", level=1)
    limitations = str(run.get("limitations") or "No additional limitations were recorded.").strip()
    for line in limitations.splitlines() or [limitations]:
        if line.strip():
            document.add_paragraph(_plain_markdown(line), style="List Bullet")

    document.add_heading("Sources and citations", level=1)
    if not citations:
        document.add_paragraph("No citations were recorded for this report.")
    for citation in citations:
        marker = str(citation.get("marker") or "Source")
        source_title = str(citation.get("source_title") or citation.get("url") or "Untitled source")
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(9)
        heading.add_run(f"{marker}  {source_title}").bold = True
        details = document.add_paragraph()
        details.add_run(f"Type: {citation.get('source_type') or 'unknown'}\n")
        if citation.get("url"):
            details.add_run(f"URL: {citation['url']}\n")
        details.add_run(f"Evidence: {citation.get('snippet') or 'No excerpt recorded.'}")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.style.font.name = "Calibri"
    footer.style.font.size = Pt(8)
    footer.style.font.color.rgb = RGBColor(102, 112, 133)
    footer_run = footer.add_run("Chronos research report  •  Page ")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(102, 112, 133)
    _add_page_field(footer)
    document.core_properties.title = str(run.get("question") or "Research report")[:255]
    document.core_properties.subject = "Chronos research export with citations and limitations"
    document.core_properties.author = "Chronos"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _pdf_text(text: str, *, strip_markdown: bool = True) -> str:
    escaped = html_escape(_plain_markdown(text) if strip_markdown else text, quote=False)
    return escaped.replace("\n", "<br/>")


def build_research_pdf(run: dict[str, Any], report: str, citations: list[dict[str, Any]]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=LETTER,
        rightMargin=0.85 * inch,
        leftMargin=0.85 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.75 * inch,
        title=str(run.get("question") or "Research report"),
        author="Chronos",
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CoverTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=28, leading=34, textColor=colors.HexColor("#1F4E79"), spaceAfter=20))
    styles.add(ParagraphStyle(name="CoverQuestion", parent=styles["BodyText"], fontName="Helvetica", fontSize=16, leading=22, textColor=colors.HexColor("#2D4059"), spaceAfter=24))
    styles.add(ParagraphStyle(name="ResearchH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=colors.HexColor("#1F4E79"), spaceBefore=14, spaceAfter=7))
    styles.add(ParagraphStyle(name="ResearchH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=17, textColor=colors.HexColor("#1F4E79"), spaceBefore=11, spaceAfter=5))
    styles.add(ParagraphStyle(name="ResearchH3", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=11.5, leading=15, textColor=colors.HexColor("#2D4059"), spaceBefore=8, spaceAfter=4))
    styles.add(ParagraphStyle(name="ResearchBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=14.5, spaceAfter=7, wordWrap="CJK"))
    styles.add(ParagraphStyle(name="ResearchBullet", parent=styles["ResearchBody"], leftIndent=16, firstLineIndent=-8, bulletIndent=4))
    styles.add(ParagraphStyle(name="ResearchCode", parent=styles["Code"], fontName="Courier", fontSize=8.5, leading=11, leftIndent=12, rightIndent=8, backColor=colors.HexColor("#F4F6F8"), borderPadding=6, spaceAfter=8, wordWrap="CJK"))
    styles.add(ParagraphStyle(name="ResearchFooter", parent=styles["BodyText"], fontName="Helvetica", fontSize=8, textColor=colors.HexColor("#667085"), alignment=TA_CENTER))

    story: list[Any] = [Spacer(1, 1.15 * inch)]
    story.append(Paragraph("Research report", styles["CoverTitle"]))
    story.append(Paragraph(_pdf_text(str(run.get("question") or "Untitled research")), styles["CoverQuestion"]))
    meta = (
        f"<b>Depth:</b> {html_escape(str(run.get('depth') or 'standard'))}<br/>"
        f"<b>Completed:</b> {html_escape(str(run.get('completed_at') or datetime.now(timezone.utc).isoformat()))}<br/>"
        f"<b>Citation policy:</b> {html_escape(str(run.get('citation_policy') or 'required'))}<br/>"
        f"<b>Sources cited:</b> {len(citations)}"
    )
    story.extend([Paragraph(meta, styles["ResearchBody"]), PageBreak()])

    style_for = {
        "heading1": styles["ResearchH1"],
        "heading2": styles["ResearchH2"],
        "heading3": styles["ResearchH3"],
    }
    for kind, raw_text in _markdown_blocks(_report_body(report)):
        if kind in style_for:
            story.append(Paragraph(_pdf_text(raw_text), style_for[kind]))
        elif kind in {"bullet", "number"}:
            bullet = "•" if kind == "bullet" else "-"
            story.append(Paragraph(_pdf_text(raw_text), styles["ResearchBullet"], bulletText=bullet))
        elif kind == "code":
            wrapped = "\n".join(
                part for line in raw_text.splitlines() for part in (textwrap.wrap(line, 92) or [""])
            )[:50_000]
            story.append(Paragraph(_pdf_text(wrapped, strip_markdown=False), styles["ResearchCode"]))
        else:
            story.append(Paragraph(_pdf_text(raw_text), styles["ResearchBody"]))

    story.append(Paragraph("Limitations", styles["ResearchH1"]))
    limitations = str(run.get("limitations") or "No additional limitations were recorded.").strip()
    for line in limitations.splitlines() or [limitations]:
        if line.strip():
            story.append(Paragraph(_pdf_text(line), styles["ResearchBullet"], bulletText="•"))

    story.append(Paragraph("Sources and citations", styles["ResearchH1"]))
    if not citations:
        story.append(Paragraph("No citations were recorded for this report.", styles["ResearchBody"]))
    for citation in citations:
        marker = str(citation.get("marker") or "Source")
        source_title = str(citation.get("source_title") or citation.get("url") or "Untitled source")
        story.append(Paragraph(f"<b>{_pdf_text(marker)} &nbsp; {_pdf_text(source_title)}</b>", styles["ResearchH2"]))
        details = [f"Type: {citation.get('source_type') or 'unknown'}"]
        if citation.get("url"):
            details.append(f"URL: {citation['url']}")
        details.append(f"Evidence: {citation.get('snippet') or 'No excerpt recorded.'}")
        story.append(Paragraph(_pdf_text("\n".join(details)), styles["ResearchBody"]))

    def footer(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#667085"))
        canvas.drawCentredString(LETTER[0] / 2, 0.38 * inch, f"Chronos research report  •  Page {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=footer, onLaterPages=footer)
    return output.getvalue()


async def _render_export_bytes(
    run: dict[str, Any],
    citations: list[dict[str, Any]],
    export_format: ResearchExportFormat,
    report_id: str,
) -> bytes:
    content = await read_artifact_content(report_id)
    if content is None:
        raise ResearchExportError("Research report content is unavailable")
    if len(content) > _MAX_REPORT_BYTES:
        raise ResearchExportError(
            f"Research report exceeds the {_MAX_REPORT_BYTES:,}-byte export limit"
        )
    if len(citations) > _MAX_CITATIONS:
        raise ResearchExportError(
            f"Research report has more than {_MAX_CITATIONS:,} citations"
        )
    citation_bytes = len(
        str(
            [
                {
                    "marker": citation.get("marker"),
                    "source_type": citation.get("source_type"),
                    "source_title": citation.get("source_title"),
                    "url": citation.get("url"),
                    "snippet": citation.get("snippet"),
                }
                for citation in citations
            ]
        ).encode("utf-8")
    )
    if citation_bytes > _MAX_CITATION_BYTES:
        raise ResearchExportError(
            f"Research citations exceed the {_MAX_CITATION_BYTES:,}-byte export limit"
        )
    try:
        report = content.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise ResearchExportError("Research report is not valid UTF-8 text") from exc
    builder = build_research_docx if export_format == "docx" else build_research_pdf
    return await asyncio.to_thread(builder, run, report, citations)


async def create_research_export(
    run: dict[str, Any],
    citations: list[dict[str, Any]],
    export_format: ResearchExportFormat,
    *,
    org_id: str,
    created_by: str,
) -> tuple[dict[str, Any], bool]:
    """Create or reuse an immutable child artifact for a completed report."""
    if export_format not in _KINDS:
        raise ResearchExportError("Export format must be docx or pdf")
    if str(run.get("status")) != "complete" or not run.get("report_artifact_id"):
        raise ResearchExportError("Only completed research reports can be exported")
    report_id = str(run["report_artifact_id"])
    report_meta = await get_artifact(report_id)
    if not report_meta or str(report_meta.get("organization_id")) != org_id or report_meta.get("is_deleted"):
        raise ResearchExportError("Research report artifact is unavailable")

    artifacts = await reflect_table("artifacts")
    existing = await _find_existing_export(
        artifacts,
        org_id=org_id,
        report_id=report_id,
        export_format=export_format,
    )
    if existing:
        if run.get("project_id") and str(existing.get("project_id") or "") != str(run["project_id"]):
            await set_artifact_project(
                str(existing["id"]),
                project_id=str(run["project_id"]),
                org_id=org_id,
            )
            existing = await get_artifact(str(existing["id"])) or existing
        return existing, True

    async with _export_lock(org_id, report_id, export_format) as lock_conn:
        # Recheck after acquiring the replica-wide lock. Every concurrent caller
        # may have rendered already, but only one is allowed to persist.
        existing = await _find_existing_export(
            artifacts,
            org_id=org_id,
            report_id=report_id,
            export_format=export_format,
            conn=lock_conn,
        )
        if existing:
            if run.get("project_id") and str(existing.get("project_id") or "") != str(run["project_id"]):
                existing = await _set_project_and_fetch(
                    artifacts,
                    str(existing["id"]),
                    project_id=str(run["project_id"]),
                    org_id=org_id,
                    conn=lock_conn,
                )
                await lock_conn.commit()
                if existing is None:
                    raise ResearchExportError("Research export could not be persisted")
            return existing, True

        # Render only after acquiring the durable cross-replica lock. A burst of
        # identical clicks therefore consumes one renderer job, not one per request.
        rendered = await _render_export_bytes(run, citations, export_format, report_id)
        artifact_id = await save_artifact(
            rendered,
            kind=_KINDS[export_format],
            title=_clean_title(str(run.get("question") or "Research report"), export_format),
            mime_type=_MIMES[export_format],
            parent_artifact_id=report_id,
            org_id=org_id,
            created_by=created_by,
            db_connection=lock_conn,
        )
        artifact = await _set_project_and_fetch(
            artifacts,
            artifact_id,
            project_id=str(run["project_id"]) if run.get("project_id") else None,
            org_id=org_id,
            conn=lock_conn,
        )
        await lock_conn.commit()
        if artifact is None:
            raise ResearchExportError("Research export could not be persisted")
        return artifact, False
